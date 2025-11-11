"""Build DOCX documents from parsed resume data."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import ParsedResume


class DocxBuilder:
    """Build DOCX documents."""

    def __init__(self):
        """Initialize builder."""
        pass

    def build_resume(self, resume: ParsedResume, output_path: Path) -> None:
        """
        Build DOCX resume from parsed data.

        Args:
            resume: Parsed resume data
            output_path: Path to save DOCX file
        """
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Determine language for headers
        language = resume.language or "en"
        headers = self._get_section_headers(language)

        # Add contact information
        self._add_contact(doc, resume.contact)

        # Add summary
        if resume.summary:
            self._add_section_header(doc, headers["summary"])
            self._add_paragraph(doc, resume.summary)

        # Add experience
        if resume.experience:
            self._add_section_header(doc, headers["experience"])
            valid_experience_count = 0
            for exp in resume.experience:
                # Only add valid experience entries (skip malformed ones)
                if exp.title and exp.title != "Unknown" and exp.title.strip():
                    # Skip if title looks like a section header (but be less aggressive)
                    title_lower = exp.title.lower().strip()
                    # Only skip if it's EXACTLY a section header word, not if it contains it
                    section_keywords = ["experience", "summary", "skills", "education", "опыт работы", "резюме", "навыки", "образование"]
                    is_section_header = title_lower in [kw.lower() for kw in section_keywords] or title_lower == "опыт"
                    
                    # Only skip very short date markers, not longer descriptions that start with "- "
                    short_markers = ["- present", "present", "-", "•"]
                    is_date_or_marker = title_lower in short_markers or (title_lower.startswith("- ") and len(title_lower) < 15)
                    
                    if not is_section_header and not is_date_or_marker:
                        # Include ALL entries that pass the basic checks - don't require bullets
                        # This ensures we include all valid experience entries even if LLM didn't generate bullets
                        self._add_experience_entry(doc, exp, language)
                        valid_experience_count += 1
            
            # If no valid experience entries were added, add a note
            if valid_experience_count == 0 and resume.experience:
                doc.add_paragraph("Experience details are being processed.")

        # Add skills
        if resume.skills:
            self._add_section_header(doc, headers["skills"])
            skills_text = ", ".join(resume.skills)
            self._add_paragraph(doc, skills_text)

        # Add education
        if resume.education:
            self._add_section_header(doc, headers["education"])
            for edu in resume.education:
                self._add_education_entry(doc, edu)

        # Save document
        doc.save(str(output_path))
    
    def _get_section_headers(self, language: str) -> dict:
        """Get section headers based on language."""
        if language == "ru":
            return {
                "summary": "Профессиональное резюме",
                "experience": "Опыт работы",
                "skills": "Навыки",
                "education": "Образование",
            }
        else:
            return {
                "summary": "Professional Summary",
                "experience": "Experience",
                "skills": "Skills",
                "education": "Education",
            }

    def build_cover_letter(
        self, cover_letter_text: str, contact_info, output_path: Path
    ) -> None:
        """
        Build DOCX cover letter.

        Args:
            cover_letter_text: Cover letter text
            contact_info: Contact information
            output_path: Path to save DOCX file
        """
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Add contact info at top
        if contact_info.name:
            para = doc.add_paragraph(contact_info.name)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if contact_info.email:
            para = doc.add_paragraph(contact_info.email)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if contact_info.phone:
            para = doc.add_paragraph(contact_info.phone)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add date
        from datetime import datetime

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.add_run(datetime.now().strftime("%B %d, %Y"))

        # Add spacing
        doc.add_paragraph()

        # Add cover letter text
        paragraphs = cover_letter_text.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                self._add_paragraph(doc, para_text.strip())

        # Save document
        doc.save(str(output_path))

    def _add_contact(self, doc: Document, contact) -> None:
        """Add contact information."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if contact.name:
            run = para.add_run(contact.name)
            run.bold = True
            run.font.size = Pt(16)

        if contact.email:
            doc.add_paragraph(contact.email).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if contact.phone:
            doc.add_paragraph(contact.phone).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if contact.location:
            doc.add_paragraph(contact.location).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

    def _add_section_header(self, doc: Document, title: str) -> None:
        """Add section header."""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        # Note: font.color is read-only in python-docx, use run.font.color.rgb instead
        try:
            run.font.color.rgb = RGBColor(0, 51, 102)
        except AttributeError:
            # Fallback if color setting fails
            pass

    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Add paragraph."""
        doc.add_paragraph(text)

    def _add_experience_entry(self, doc: Document, exp, language: str = "en") -> None:
        """Add experience entry."""
        # Title and company
        para = doc.add_paragraph()
        run = para.add_run(exp.title)
        run.bold = True
        if exp.company and exp.company != "Unknown" and exp.company.strip():
            separator = " в " if language == "ru" else " at "
            para.add_run(f"{separator}{exp.company}")

        # Dates
        if exp.dates and exp.dates.strip():
            para = doc.add_paragraph(exp.dates)
            para.italic = True

        # Bullets - ensure we add them even if empty, but try to preserve original if LLM failed
        if exp.bullets:
            for bullet in exp.bullets:
                if bullet and bullet.strip():
                    # Clean up bullet text
                    bullet_clean = bullet.strip()
                    # Remove leading dashes/bullets if present
                    bullet_clean = bullet_clean.lstrip("-•* ").strip()
                    if bullet_clean:
                        para = doc.add_paragraph(bullet_clean, style="List Bullet")
        else:
            # If no bullets, add a placeholder or use raw_text if available
            if exp.raw_text and exp.raw_text.strip():
                # Try to extract meaningful content from raw_text
                raw_lines = [line.strip() for line in exp.raw_text.split("\n") if line.strip()]
                for line in raw_lines[:5]:  # Limit to 5 lines
                    if len(line) > 10 and not line.lower().startswith(("title:", "company:", "dates:")):
                        para = doc.add_paragraph(line, style="List Bullet")

        doc.add_paragraph()  # Spacing

    def _add_education_entry(self, doc: Document, edu) -> None:
        """Add education entry."""
        para = doc.add_paragraph()
        if edu.degree:
            run = para.add_run(edu.degree)
            run.bold = True
            if edu.institution:
                para.add_run(f", {edu.institution}")
        elif edu.institution:
            para.add_run(edu.institution).bold = True

        if edu.dates:
            doc.add_paragraph(edu.dates).italic = True

        if edu.details:
            doc.add_paragraph(edu.details)

        doc.add_paragraph()  # Spacing

